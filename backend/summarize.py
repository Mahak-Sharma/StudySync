import os
import pytesseract
from PIL import Image
import fitz  # PyMuPDF
import docx
from transformers import pipeline
import io

summarizer = pipeline("summarization", model="sshleifer/distilbart-cnn-12-6")

def extract_text_from_pdf(file_obj):
    """Extract text from PDF file object without saving to disk"""
    text = ""
    # Read the file object into memory
    pdf_data = file_obj.read()
    file_obj.seek(0)  # Reset file pointer for potential future reads
    
    # Open PDF from memory
    with fitz.open(stream=pdf_data, filetype="pdf") as doc:
        for page in doc:
            text += page.get_text()
    return text

def extract_text_from_docx(file_obj):
    """Extract text from DOCX file object without saving to disk"""
    # Read the file object into memory
    docx_data = file_obj.read()
    file_obj.seek(0)  # Reset file pointer for potential future reads
    
    # Open DOCX from memory
    doc = docx.Document(io.BytesIO(docx_data))
    return "\n".join([para.text for para in doc.paragraphs])

def extract_text_from_image(file_obj):
    """Extract text from image file object without saving to disk"""
    # Read the file object into memory
    image_data = file_obj.read()
    file_obj.seek(0)  # Reset file pointer for potential future reads
    
    # Open image from memory
    image = Image.open(io.BytesIO(image_data))
    return pytesseract.image_to_string(image)

def summarize_text(text, max_len=1300):
    if not text.strip():
        return "No text found for summarization."
    
    if len(text) > 1000:
        text = text[:1000] 
    
    summary = summarizer(text, max_length=130, min_length=30, do_sample=False)
    return summary[0]['summary_text']

# Keep the old file_path functions for backward compatibility
def extract_text_from_pdf_file(file_path):
    """Legacy function for file paths - kept for compatibility"""
    text = ""
    with fitz.open(file_path) as doc:
        for page in doc:
            text += page.get_text()
    return text

def extract_text_from_docx_file(file_path):
    """Legacy function for file paths - kept for compatibility"""
    doc = docx.Document(file_path)
    return "\n".join([para.text for para in doc.paragraphs])

def extract_text_from_image_file(file_path):
    """Legacy function for file paths - kept for compatibility"""
    image = Image.open(file_path)
    return pytesseract.image_to_string(image)

def main():
    file_path = input("Enter the file path (.pdf, .docx, image): ").strip()

    if not os.path.exists(file_path):
        print("File not found.")
        return

    ext = os.path.splitext(file_path)[1].lower()
    
    if ext == ".pdf":
        text = extract_text_from_pdf_file(file_path)
    elif ext == ".docx":
        text = extract_text_from_docx_file(file_path)
    elif ext in [".png", ".jpg", ".jpeg", ".bmp", ".tiff"]:
        text = extract_text_from_image_file(file_path)
    else:
        print("Unsupported file format.")
        return

    print("\nExtracted Text:\n", text[:1000] + "..." if len(text) > 1000 else text)

    summary = summarize_text(text)
    print("\n--- Summary ---\n", summary)

if __name__ == "__main__":
    main()
